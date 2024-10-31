from tkinter import filedialog
import pandas as pd
import time
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from tqdm import tqdm

from tkinter import Tk, simpledialog
from tkinter.filedialog import askopenfilename

from pathlib import Path


def main():

    # Disable Tkinter main interface
    Tk().withdraw()

    # Open file dialog to select the Excel file
    excel_file_path = askopenfilename(
        title="Select Excel File", filetypes=[("Excel files", "*.xlsx *.xls")]
    )

    # Full path
    full_path = Path(excel_file_path)

    # Separate the path and the filename
    path = full_path.parent
    filename = full_path.name

    print("Path:", path)
    print("Filename:", filename)

    # Check if a file has been selected
    if excel_file_path:
        # Load the selected Excel file
        xls = pd.ExcelFile(excel_file_path)

        # Display available sheet names
        sheet_names = xls.sheet_names
        print("\nAvailable sheets in this Excel file:")
        for idx, sheet_name in enumerate(sheet_names):
            print(f"{idx + 1}. {sheet_name}")

        # Ask the user to select which sheet to use
        while True:
            try:
                choice = int(input("\nEnter the number of the sheet to use: ")) - 1
                if 0 <= choice < len(sheet_names):
                    selected_sheet = sheet_names[choice]
                    break
                else:
                    print("Please enter a valid number.")
            except ValueError:
                print("Please enter a valid number.")

        # Load the selected sheet
        df = pd.read_excel(xls, sheet_name=selected_sheet)

        # Display the first few rows to verify the file has been loaded correctly
        print(f"\nLoading sheet: {selected_sheet}")
        print(df.head())
    else:
        print("No file was selected.")

    # Ask the user to select which columns to inspect
    print("\nAvailable columns in the selected sheet:")
    for idx, col_name in enumerate(df.columns):
        print(f"{idx + 1}. {col_name}")

    selected_columns = []
    while True:
        try:
            col_choice = input(
                "\nEnter the numbers of the columns to inspect (comma-separated, or press Enter to finish): "
            )
            if not col_choice:
                break
            col_indices = [int(num.strip()) - 1 for num in col_choice.split(",")]
            for idx in col_indices:
                if 0 <= idx < len(df.columns):
                    selected_columns.append(df.columns[idx])
                else:
                    print(
                        f"Column index {idx + 1} is out of range. Please enter a valid number."
                    )
            if selected_columns:
                break
        except ValueError:
            print("Please enter valid numbers separated by commas.")

    print(f"\nSelected columns to inspect: {selected_columns}")

    # Ask the user to select which columns to keep in the output file
    selected_output_columns = []
    while True:
        try:
            output_col_choice = input(
                "\nEnter the numbers of the columns to keep in the output file (comma-separated, or press Enter to finish): "
            )
            if not output_col_choice:
                break
            output_col_indices = [
                int(num.strip()) - 1 for num in output_col_choice.split(",")
            ]
            for idx in output_col_indices:
                if 0 <= idx < len(df.columns):
                    selected_output_columns.append(df.columns[idx])
                else:
                    print(
                        f"Column index {idx + 1} is out of range. Please enter a valid number."
                    )
            if selected_output_columns:
                break
        except ValueError:
            print("Please enter valid numbers separated by commas.")

    print(f"\nSelected columns to keep in output: {selected_output_columns}")

    # Ask the user for categories and corresponding keywords to highlight
    highlight_categories = {}
    while True:
        category_name = simpledialog.askstring(
            "Input Category", "Enter a category name (or press Cancel to finish):"
        )
        if not category_name:
            break
        keywords_input = simpledialog.askstring(
            "Input Keywords",
            f"Enter keywords for category '{category_name}' (comma-separated):",
        )
        if keywords_input:
            highlight_categories[category_name] = [
                kw.strip() for kw in keywords_input.split(",")
            ]

    # Define a palette of highlight colors that are colorblind-friendly and not too light
    highlight_colors = [
        WD_COLOR_INDEX.YELLOW,  # Yellow (distinct and visible)
        WD_COLOR_INDEX.BRIGHT_GREEN,  # Bright Green (distinct and visible)
        WD_COLOR_INDEX.RED,  # Red (distinct)
        WD_COLOR_INDEX.BLUE,  # Blue (distinct and visible)
        WD_COLOR_INDEX.GRAY_50,  # Dark Gray (distinct and visible)
        WD_COLOR_INDEX.DARK_YELLOW,  # Dark Yellow (good contrast)
        WD_COLOR_INDEX.DARK_RED,  # Dark Red (good contrast)
    ]

    # Assign colors automatically to categories
    category_colors = {}
    for i, category in enumerate(highlight_categories):
        category_colors[category] = highlight_colors[i % len(highlight_colors)]

    # Display summary before processing
    total_rows = len(df)
    output_word_structured_path = f"{path}/highlighted_document.docx"
    # Specify the output document path
    output_word_structured_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        title="Save Word file as",
        filetypes=[("Word files", "*.docx")],
    )

    print(f"\nProcessing file: {filename}")
    print("\nCategories and keywords to highlight:")
    for category, keywords in highlight_categories.items():
        print(f"- {category}: {', '.join(keywords)}")
    print(f"\nTotal rows to process: {total_rows}")
    print(f"Output file will be saved as: {output_word_structured_path}\n")

    # Create the Word document
    doc = Document()
    doc.add_heading("Highlighted Articles - Structured View", level=1)

    # Using tqdm for the progress bar
    for index, row in tqdm(df.iterrows(), total=total_rows, desc="Processing rows"):
        # Add a clear separation between each article
        doc.add_paragraph("\n" + "=" * 50 + "\n")

        # Write only selected output columns
        for col_name in selected_output_columns:
            if col_name in row and pd.notna(row[col_name]):
                paragraph = doc.add_paragraph(f"{col_name}: ")
                cell_value = str(row[col_name])
                if col_name in selected_columns:  # Highlight specific keywords
                    for word in cell_value.split():
                        run = paragraph.add_run(word + " ")
                        for category, keywords in highlight_categories.items():
                            if any(
                                keyword.lower() in word.lower() for keyword in keywords
                            ):
                                run.font.highlight_color = category_colors[category]
                                break
                else:
                    paragraph.add_run(cell_value)

    # Save the Word document
    doc.save(output_word_structured_path)

    # Print the final file path
    print(f"\nProcessing complete. File saved as: {output_word_structured_path}")


if __name__ == "__main__":
    main()
